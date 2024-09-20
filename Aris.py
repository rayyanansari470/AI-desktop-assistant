from AppOpener import open, close
from docx import Document
import pyttsx3
import speech_recognition as sr
import os
import time
import webbrowser
import cohere
import keyboard
import shutil

# Initialize the assistant status
is_active = True

# Initialize Cohere API
co = cohere.Client('YOUR-COHERE-API-KEY')

# Using time and pyttsx3 module for greetings
engine = pyttsx3.init()
voice_id = 'HKEY_LOCAL_MACHINE\\SOFTWARE\\Microsoft\\Speech\\Voices\\Tokens\\TTS_MS_EN-US_ZIRA_11.0'
engine.setProperty('voice', voice_id)

current_hour = int(time.strftime('%H'))
if current_hour < 12:
    print("Good morning Rayyan, I am Aris, how can I help you?")
    engine.say("Good morning Rayyan, I am Aris, how can I help you?")
elif 12 <= current_hour < 16:
    print("Good afternoon Rayyan, I am Aris, how can I help you?")
    engine.say("Good afternoon Rayyan, I am Aris, how can I help you?")
else:
    print("Good evening Rayyan, I am Aris, how can I help you?")
    engine.say("Good evening Rayyan, I am Aris, how can I help you?")
engine.runAndWait()

# Storing recognizer function in recognizer variable
recognizer = sr.Recognizer()

chat = ""


def talk(cmd):
    global chat

    # Setting up prompt
    chat += f"Rayyan: {cmd}\nAris:"

    # Cohere API call
    response = co.generate(
        model='command-xlarge-nightly',
        prompt=chat,
        max_tokens=150,
        temperature=0.7
    )

    reply = response.generations[0].text[5:].strip()

    # Speak the response
    engine.say(reply)
    engine.runAndWait()

    # Append the response to chat for continuity
    chat += f"{reply}\n"
    return reply


def ai(prompt):
    text = f"Cohere response for Prompt: {prompt} \n**********************************\n\n"

    # Cohere API call for the AI response
    response = co.generate(
        model='command-xlarge-nightly',
        prompt=prompt,
        max_tokens=150,
        temperature=0.7
    )

    reply = response.generations[0].text[5:].strip()

    # Print the response
    print(reply)
    text += reply

    # Create a .docx file
    if not os.path.exists("cohere_responses"):
        os.mkdir("cohere_responses")

    doc = Document()
    doc.add_heading('AI Response', 0)
    doc.add_paragraph(text)

     # Save the document in the current directory first
    file_path = "cohere_responses/prompt_response.docx"
    doc.save(file_path)
    print("Document created successfully!")

    # Get the desktop path
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")

    # Move the file to the desktop
    try:
        shutil.move(file_path, desktop_path)
        print(f"Document moved to {desktop_path} successfully!")
    except Exception as e:
        print(f"Error moving the file: {e}")


# Function for listening to commands
def listen_for_command():
    with sr.Microphone() as source:  # Using microphone named as source
        print("Listening...")
        recognizer.adjust_for_ambient_noise(source)  # Adjusts for ambient noise
        audio = recognizer.listen(source)  # Listens to the voice
    try:
        print("Recognizing...")
        cmd = recognizer.recognize_google(audio, language="en-in")  # Recognizes the voice
        print("User:", cmd)
        return cmd.lower()
    except sr.UnknownValueError:
        print("Sorry, I didn't catch that")
        engine.say("Sorry, I didn't catch that")
        engine.runAndWait()
        return ""
    except sr.RequestError:
        print("Could not request results. Please check your internet connection")
        engine.say("Could not request results. Please check your internet connection")
        engine.runAndWait()
        return ""
    
# Function for executing command
def execute_command(cmd):
    global is_active

    if "turn off please" in cmd:
        is_active = False
        engine.say("Goodbye!")
        engine.runAndWait()
        return

    if is_active:
        if "open google" in cmd:
            webbrowser.open("https://www.google.com")
        elif "close browser" in cmd:
            os.system("taskkill /f /im chrome.exe")
        elif "open youtube" in cmd:
            webbrowser.open("https://www.youtube.com")
        elif "open whatsapp" in cmd:
            open("Whatsapp")
        elif "close whatsapp" in cmd:
            close("Whatsapp")
        elif "open notepad" in cmd:
            open("notepad")
        elif "close notepad" in cmd:
            close("notepad")
        elif "shutdown the computer" in cmd:
            os.system("shutdown /s /t 1")
        elif "meet my family" in cmd:
            engine.say("Hello dear siblings! Nice to meet you all. I am Aris")
            engine.runAndWait()
        elif "what is current time" in cmd:
            current_time = time.localtime(time.time())
            say_time = "Current time is " + time.asctime(current_time)
            engine.say(say_time)
            engine.runAndWait()
        elif "search for" in cmd:
            search_query = cmd.split("search for ")[1].strip()
            search_url = "https://www.google.com/search?q=" + search_query.replace(" ", "+")
            webbrowser.open(search_url)
        elif "using artificial intelligence" in cmd:
            ai(prompt=cmd)
        else:
            talk(cmd)

def main():
    global is_active

    while True:
        # Check for manual keyboard control
        if keyboard.is_pressed('ctrl+shift+o'):  # Turn on manually
            if not is_active:
                is_active = True
                engine.say("Yes? How can I help you?")
                engine.runAndWait()

        # Listen for voice command only when active
        if is_active:
            command = listen_for_command()
            if command:
                execute_command(command)
            
        time.sleep(1)


if __name__ == "__main__":
    main()
