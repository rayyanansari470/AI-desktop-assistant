from AppOpener import open, close
from docx import Document
import pyttsx3
import speech_recognition as sr
import os
import time
import webbrowser
import cohere
import keyboard
import shutil

# Initialize the assistant status
is_active = True

# Initialize Cohere API
co = cohere.Client('YOUR-COHERE-API-KEY')

# Using time and pyttsx3 module for greetings
engine = pyttsx3.init()
voice_id = 'HKEY_LOCAL_MACHINE\\SOFTWARE\\Microsoft\\Speech\\Voices\\Tokens\\TTS_MS_EN-US_ZIRA_11.0'
engine.setProperty('voice', voice_id)

current_hour = int(time.strftime('%H'))
if current_hour < 12:
    print("Good morning Rayyan, I am Luna, how can I help you?")
    engine.say("Good morning Rayyan, I am Luna, how can I help you?")
elif 12 <= current_hour < 16:
    print("Good afternoon Rayyan, I am Luna, how can I help you?")
    engine.say("Good afternoon Rayyan, I am Luna, how can I help you?")
else:
    print("Good evening Rayyan, I am Luna, how can I help you?")
    engine.say("Good evening Rayyan, I am Luna, how can I help you?")
engine.runAndWait()

# Storing recognizer function in recognizer variable
recognizer = sr.Recognizer()

chat = ""


def talk(cmd):
    global chat

    # Setting up prompt
    chat += f"Rayyan: {cmd}\nLuna:"

    # Cohere API call
    response = co.generate(
        model='command-xlarge-nightly',
        prompt=chat,
        max_tokens=150,
        temperature=0.7
    )

    reply = response.generations[0].text[5:].strip()

    # Speak the response
    engine.say(reply)
    engine.runAndWait()

    # Append the response to chat for continuity
    chat += f"{reply}\n"
    return reply


def ai(prompt):
    text = f"Cohere response for Prompt: {prompt} \n**********************************\n\n"

    # Cohere API call for the AI response
    response = co.generate(
        model='command-xlarge-nightly',
        prompt=prompt,
        max_tokens=150,
        temperature=0.7
    )

    reply = response.generations[0].text[5:].strip()

    # Print the response
    print(reply)
    text += reply

    # Create a .docx file
    if not os.path.exists("cohere_responses"):
        os.mkdir("cohere_responses")

    doc = Document()
    doc.add_heading('AI Response', 0)
    doc.add_paragraph(text)

     # Save the document in the current directory first
    file_path = "cohere_responses/prompt_response.docx"
    doc.save(file_path)
    print("Document created successfully!")

    # Get the desktop path
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")

    # Move the file to the desktop
    try:
        shutil.move(file_path, desktop_path)
        print("Document moved to desktop successfully!")
        engine.say("Document moved to desktop successfully!")
        engine.runAndWait()
    except Exception as e:
        print(f"Error moving the file: {e}")


# Function for listening to commands
def listen_for_command():
    with sr.Microphone() as source:  # Using microphone named as source
        print("Listening...")
        recognizer.adjust_for_ambient_noise(source)  # Adjusts for ambient noise
        audio = recognizer.listen(source)  # Listens to the voice
    try:
        print("Recognizing...")
        cmd = recognizer.recognize_google(audio, language="en-in")  # Recognizes the voice
        print("User:", cmd)
        return cmd.lower()
    except sr.UnknownValueError:
        print("Sorry, I didn't catch that")
        engine.say("Sorry, I didn't catch that")
        engine.runAndWait()
        return ""
    except sr.RequestError:
        print("Could not request results. Please check your internet connection")
        engine.say("Could not request results. Please check your internet connection")
        engine.runAndWait()
        return ""
    
# Function for executing command
def execute_command(cmd):
    global is_active

    if "turn off please" in cmd:
        is_active = False
        engine.say("Goodbye!")
        engine.runAndWait()
        return

    if is_active:
        sites = [["google","https://www.google.com" ], ["youtube", "https://www.youtube.com"], ["wikipedia", "https://www.wikipedia.com"]]
        for site in sites:
            if f"open {site[0]}" in cmd:
                engine.say(f"opening {site[0]}!")
                engine.runAndWait()
                webbrowser.open(site[1])
        if "close browser" in cmd:
            engine.say("closing browser!")
            engine.runAndWait()
            os.system("taskkill /f /im chrome.exe")
        elif "open whatsapp" in cmd:
            engine.say("opening whatsapp!")
            engine.runAndWait()
            open("Whatsapp")
        elif "close whatsapp" in cmd:
            engine.say("closing whatsapp!")
            engine.runAndWait()
            close("Whatsapp")
        elif "open notepad" in cmd:
            engine.say("opening notepad!")
            engine.runAndWait()
            open("notepad")
        elif "close notepad" in cmd:
            engine.say("closing notepad!")
            engine.runAndWait()
            close("notepad")
        elif "shutdown the computer" in cmd:
            engine.say("shutting down!")
            engine.runAndWait()
            os.system("shutdown /s /t 1")
        elif "what is current time" in cmd:
            current_time = time.localtime(time.time())
            say_time = "Current time is " + time.asctime(current_time)
            engine.say(say_time)
            engine.runAndWait()
        elif "search for" in cmd:
            search_query = cmd.split("search for ")[1].strip()
            search_url = "https://www.google.com/search?q=" + search_query.replace(" ", "+")
            engine.say("searching!")
            engine.runAndWait()
            webbrowser.open(search_url)
        elif "using artificial intelligence" in cmd:
            ai(prompt=cmd)
        else:
            talk(cmd)

def main():
    global is_active

    while True:
        # Check for manual keyboard control
        if keyboard.is_pressed('ctrl+shift+o'):  # Turn on manually
            if not is_active:
                is_active = True
                engine.say("Yes Rayyan? How can I help you?")
                engine.runAndWait()

        # Listen for voice command only when active
        if is_active:
            command = listen_for_command()
            if command:
                execute_command(command)
            
        time.sleep(1)


if __name__ == "__main__":
    main()
